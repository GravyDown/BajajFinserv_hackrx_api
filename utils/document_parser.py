import os
import re
import tempfile
import requests
from typing import List, Tuple, Optional
from urllib.parse import urlparse, parse_qs
import PyPDF2
from docx import Document
import email
from email import policy
from email.parser import BytesParser
from bs4 import BeautifulSoup
import chardet
from tenacity import retry, stop_after_attempt, wait_exponential

from models.schemas import DocumentType

class DocumentParser:
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
        
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    def download_file(self, url: str) -> Tuple[str, DocumentType]:
        """Download file from URL (including Google Drive support)"""
        # Handle Google Drive URLs
        if 'drive.google.com' in url:
            url = self._convert_drive_url(url)
            
        response = requests.get(url, stream=True, timeout=30)
        response.raise_for_status()
        
        # Detect file type
        content_type = response.headers.get('content-type', '').lower()
        filename = self._extract_filename(response.headers.get('content-disposition', ''))
        
        if not filename:
            filename = os.path.basename(urlparse(url).path) or 'document'
            
        doc_type = self._detect_document_type(content_type, filename)
        
        # Save to temp file
        temp_path = os.path.join(self.temp_dir, filename)
        with open(temp_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
                
        return temp_path, doc_type
    
    def _convert_drive_url(self, url: str) -> str:
        """Convert Google Drive URL to direct download link"""
        # Extract file ID from various Google Drive URL formats
        file_id = None
        
        if '/file/d/' in url:
            file_id = url.split('/file/d/')[1].split('/')[0]
        elif 'id=' in url:
            parsed = urlparse(url)
            file_id = parse_qs(parsed.query).get('id', [None])[0]
            
        if file_id:
            return f"https://drive.google.com/uc?export=download&id={file_id}"
        return url
    
    def _extract_filename(self, content_disposition: str) -> Optional[str]:
        """Extract filename from Content-Disposition header"""
        if not content_disposition:
            return None
            
        filename_match = re.search(r'filename="?([^"]+)"?', content_disposition)
        if filename_match:
            return filename_match.group(1)
        return None
    
    def _detect_document_type(self, content_type: str, filename: str) -> DocumentType:
        """Detect document type from content type and filename"""
        filename_lower = filename.lower()
        
        if 'pdf' in content_type or filename_lower.endswith('.pdf'):
            return DocumentType.PDF
        elif 'wordprocessingml' in content_type or filename_lower.endswith(('.docx', '.doc')):
            return DocumentType.DOCX
        elif 'message/rfc822' in content_type or filename_lower.endswith('.eml'):
            return DocumentType.EMAIL
        else:
            return DocumentType.UNKNOWN
    
    def parse_document(self, file_path: str, doc_type: DocumentType) -> str:
        """Parse document based on type"""
        if doc_type == DocumentType.PDF:
            return self._parse_pdf(file_path)
        elif doc_type == DocumentType.DOCX:
            return self._parse_docx(file_path)
        elif doc_type == DocumentType.EMAIL:
            return self._parse_email(file_path)
        else:
            # Try to detect encoding and read as text
            return self._parse_text(file_path)
    
    def _parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
        return text.strip()
    
    def _parse_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
                        
            return text.strip()
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    def _parse_email(self, file_path: str) -> str:
        """Extract text from email file"""
        try:
            with open(file_path, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)
                
            text = f"Subject: {msg['subject']}\n"
            text += f"From: {msg['from']}\n"
            text += f"To: {msg['to']}\n"
            text += f"Date: {msg['date']}\n\n"
            
            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == 'text/plain':
                        text += part.get_content() + "\n"
                    elif part.get_content_type() == 'text/html':
                        html_content = part.get_content()
                        soup = BeautifulSoup(html_content, 'html.parser')
                        text += soup.get_text() + "\n"
            else:
                text += msg.get_content()
                
            return text.strip()
        except Exception as e:
            raise Exception(f"Error parsing email: {str(e)}")
    
    def _parse_text(self, file_path: str) -> str:
        """Parse generic text file with encoding detection"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
                
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read().strip()
        except Exception as e:
            raise Exception(f"Error parsing text file: {str(e)}")
    
    def cleanup(self):
        """Clean up temporary files"""
        import shutil
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)